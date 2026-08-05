from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from bibagent.models import ProjectPaths
from bibagent.word_export import (
    DEFAULT_FIGURES,
    FIGURE_RESULT_SECTION_PREFIXES,
    export_word_report,
)


def test_word_export_places_each_figure_once_in_results(tmp_path: Path) -> None:
    paths = ProjectPaths(tmp_path)
    paths.create()
    lines = [
        "# 测试文献计量论文",
        "## 摘要",
        "**目的**：测试Word导出。",
        "## 2 数据与方法",
        "语料库包含269篇唯一文献，但方法部分不应插图。",
        "## 3 结果",
    ]
    for section_prefix in ("3.1", "3.2", "3.3", "3.4", "3.5"):
        lines.append(f"### {section_prefix} 测试结果")
        figure_numbers = [
            number
            for number, prefix in FIGURE_RESULT_SECTION_PREFIXES.items()
            if prefix == section_prefix
        ]
        lines.append(
            "本小节依次报告"
            + "、".join(f"图{number}" for number in figure_numbers)
            + "，这是结果段落[E001]。"
        )
    lines.extend(
        [
            "## 图表解读清单",
            "| 图表编号 | 名称 |",
            "|---|---|",
            "| 图1 | 测试 |",
            "## 参考文献",
            "1. Example reference.",
        ]
    )
    (paths.report / "manuscript.md").write_text("\n\n".join(lines), encoding="utf-8")
    for index, figure in enumerate(DEFAULT_FIGURES):
        Image.new("RGB", (1200, 800), (35 + index * 5, 90, 160)).save(
            paths.figures / figure.filename
        )

    output = export_word_report(tmp_path)
    document = Document(output)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    captions = [text for text in paragraphs if text.startswith("图 ")]
    reading_guides = [text for text in paragraphs if "直观解读。" in text]

    assert len(document.inline_shapes) == 16
    assert len(captions) == 16
    assert len(reading_guides) == 16
    assert captions[0] == "图 1  年度科学产出趋势"
    assert captions[-1] == "图 16  文献耦合网络"
    assert [int(text.split()[1]) for text in captions] == list(range(1, 17))
    assert not any("图表解读清单" in text for text in paragraphs)
    assert not any("[E001]" in text for text in paragraphs)
    section = document.sections[0]
    assert round(section.page_width.cm, 1) == 21.0
    assert round(section.page_height.cm, 1) == 29.7
    heading_fonts = (
        document.styles["Heading 1"]
        .element.get_or_add_rPr()
        .get_or_add_rFonts()
    )
    assert heading_fonts.get(qn("w:eastAsia")) == "Microsoft YaHei"
    title_run_fonts = (
        document.paragraphs[0].runs[0]
        ._element.get_or_add_rPr()
        .get_or_add_rFonts()
    )
    assert title_run_fonts.get(qn("w:eastAsia")) == "Microsoft YaHei"

    for caption in captions:
        caption_index = paragraphs.index(caption)
        next_caption = next(
            (
                index
                for index in range(caption_index + 1, len(paragraphs))
                if paragraphs[index].startswith("图 ")
            ),
            len(paragraphs),
        )
        assert any(
            "直观解读。" in paragraphs[index]
            for index in range(caption_index + 1, next_caption)
        )


def test_cross_reference_does_not_move_figure_to_wrong_subsection(tmp_path: Path) -> None:
    paths = ProjectPaths(tmp_path)
    paths.create()
    lines = [
        "# 测试文献计量论文",
        "## 3 结果",
        "### 3.1 绩效",
        "图1、图2、图3和图4展示绩效；图14在此仅用于跨结果比较，不应在本小节插入。",
        "### 3.2 社会结构",
        "图5、图6、图7和图8展示社会结构。",
        "### 3.3 引文影响",
        "图9和图10展示引文影响。",
        "### 3.4 概念结构",
        "图11、图12、图13和图14展示概念结构。",
        "### 3.5 知识基础",
        "图15和图16展示知识基础。",
    ]
    (paths.report / "manuscript.md").write_text("\n\n".join(lines), encoding="utf-8")
    for index, figure in enumerate(DEFAULT_FIGURES):
        Image.new("RGB", (1200, 800), (35 + index * 5, 90, 160)).save(
            paths.figures / figure.filename
        )

    output = export_word_report(tmp_path)
    paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]

    conceptual_heading = paragraphs.index("3.4 概念结构")
    figure_14_caption = paragraphs.index("图 14  作者—来源—关键词三字段关系")
    assert figure_14_caption > conceptual_heading


def test_word_export_converts_internal_evidence_to_numbered_citations(
    tmp_path: Path,
) -> None:
    paths = ProjectPaths(tmp_path)
    paths.create()
    lines = [
        "# 测试文献计量论文",
        "## 摘要",
        "摘要中的内部证据不应显示[E101][E102]。",
        "## 1 引言",
        "第一处外部引文[E102]，实证证据不显示[E001]，连续引文应压缩[E102][E101][E103]。",
        "## 3 结果",
    ]
    for section_prefix in ("3.1", "3.2", "3.3", "3.4", "3.5"):
        lines.extend(
            [
                f"### {section_prefix} 测试结果",
                "结果由内部证据支持[E001]。",
            ]
        )
    lines.extend(
        [
            "## 参考文献",
            "1. Reference A. https://doi.org/10.1/a [E101]",
            "2. Reference B. https://doi.org/10.1/b [E102]",
            "3. Reference C. https://doi.org/10.1/c [E103]",
        ]
    )
    (paths.report / "manuscript.md").write_text("\n\n".join(lines), encoding="utf-8")
    for index, figure in enumerate(DEFAULT_FIGURES):
        Image.new("RGB", (1200, 800), (35 + index * 5, 90, 160)).save(
            paths.figures / figure.filename
        )

    document = Document(export_word_report(tmp_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    references = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Reference Entry"
    ]

    assert "[E" not in text
    assert "第一处外部引文[1]" in text
    assert "连续引文应压缩[1–3]" in text
    assert references == [
        "[1] Reference B. https://doi.org/10.1/b",
        "[2] Reference A. https://doi.org/10.1/a",
        "[3] Reference C. https://doi.org/10.1/c",
    ]
