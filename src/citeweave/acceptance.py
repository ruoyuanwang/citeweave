from __future__ import annotations

import gzip
import hashlib
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import networkx as nx
import pandas as pd
from docx import Document
from PIL import Image

from .generation import SECTION_CONTRACTS, validate_manuscript
from .io import read_json, sha256_file, write_json
from .models import AcquisitionManifest, ProjectPaths
from .workflow import _load_evidence, _package_manifest

REQUIRED_TABLES = {
    "works": {"work_id", "title", "year", "doi", "source_id", "source_record_hash"},
    "authors": {"author_id", "name"},
    "institutions": {"institution_id", "name"},
    "authorships": {"work_id", "author_id", "institution_id"},
    "sources": {"source_id", "name"},
    "keywords": {"work_id", "keyword"},
    "topics": {"work_id"},
    "references": {"citing_work_id", "cited_work_id"},
    "provenance": {"work_id", "source", "source_record_hash"},
    "duplicates": {"removed_work_id", "kept_work_id", "rule"},
}


def _check(name: str, passed: bool, detail: Any) -> dict[str, Any]:
    return {"name": name, "passed": bool(passed), "detail": detail}


def _raw_payload_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    opener = gzip.open if path.suffix == ".gz" else path.open
    with opener(path, "rb") if path.suffix == ".gz" else opener("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def verify_project(
    root: Path,
    *,
    require_llm_model: str | None = None,
    minimum_figures: int = 12,
) -> dict[str, Any]:
    """Run deterministic acceptance checks against an existing research package."""
    paths = ProjectPaths(root)
    checks: list[dict[str, Any]] = []
    manifest = AcquisitionManifest.model_validate(
        read_json(paths.audit / "acquisition_manifest.json")
    )
    exact = (
        manifest.complete
        and not manifest.truncated
        and manifest.drift == 0
        and (
            manifest.expected_records is None
            or manifest.unique_records == manifest.expected_records
        )
    )
    checks.append(
        _check(
            "source_relative_acquisition_completeness",
            exact,
            {
                "source": manifest.source.value,
                "expected": manifest.expected_records,
                "received": manifest.received_records,
                "unique": manifest.unique_records,
                "drift": manifest.drift,
                "truncated": manifest.truncated,
            },
        )
    )

    raw_hashes = {
        _raw_payload_sha256(path) for path in paths.raw.rglob("*") if path.is_file()
    }
    missing_hashes = sorted(set(manifest.raw_sha256) - raw_hashes)
    checks.append(
        _check(
            "raw_snapshot_hashes",
            not missing_hashes and bool(manifest.raw_sha256),
            {"declared": len(manifest.raw_sha256), "missing": missing_hashes},
        )
    )

    table_rows: dict[str, int] = {}
    schema_errors: dict[str, list[str]] = {}
    for name, required_columns in REQUIRED_TABLES.items():
        path = paths.canonical / f"{name}.parquet"
        if not path.exists():
            schema_errors[name] = ["<file missing>"]
            continue
        frame = pd.read_parquet(path)
        table_rows[name] = len(frame)
        missing = sorted(required_columns - set(frame.columns))
        if missing:
            schema_errors[name] = missing
    checks.append(
        _check(
            "canonical_star_schema",
            not schema_errors and table_rows.get("works", 0) > 0,
            {"rows": table_rows, "schema_errors": schema_errors},
        )
    )

    quality = read_json(paths.quality / "summary.json")
    quality_pass = (
        quality.get("acquisition_complete") is True
        and quality.get("truncated") is False
        and quality.get("canonical_works") == table_rows.get("works")
    )
    checks.append(_check("quality_gate", quality_pass, quality))

    figure_manifest = read_json(paths.figures / "figure_manifest.json")
    figures = (
        figure_manifest.get("figures", [])
        if isinstance(figure_manifest, dict)
        else figure_manifest
    )
    figure_errors = []
    for figure in figures:
        png_value, svg_value = Path(figure["png"]), Path(figure["svg"])
        png = png_value if png_value.is_absolute() else paths.figures / png_value
        svg = svg_value if svg_value.is_absolute() else paths.figures / svg_value
        if not png.exists() or not svg.exists():
            figure_errors.append(f"{figure['name']}: missing output")
            continue
        with Image.open(png) as image:
            width, height = image.size
        if min(width, height) < 1000:
            figure_errors.append(f"{figure['name']}: only {width}x{height}px")
        qa = figure.get("qa", figure)
        if sha256_file(png) != qa.get("png_sha256"):
            figure_errors.append(f"{figure['name']}: PNG hash mismatch")
        if sha256_file(svg) != qa.get("svg_sha256"):
            figure_errors.append(f"{figure['name']}: SVG hash mismatch")
    checks.append(
        _check(
            "publication_figures",
            len(figures) >= minimum_figures and not figure_errors,
            {"count": len(figures), "minimum": minimum_figures, "errors": figure_errors},
        )
    )

    word_candidates = [
        paths.report / "manuscript.docx",
        paths.report / "bibliometric_report.docx",
    ]
    word_path = next((path for path in word_candidates if path.exists()), None)
    word_detail: dict[str, Any] = {"path": None}
    word_pass = False
    if word_path is not None:
        word_document = Document(word_path)
        captions = [
            paragraph.text
            for paragraph in word_document.paragraphs
            if paragraph.text.startswith("图 ")
        ]
        word_text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
        internal_evidence_tokens = re.findall(r"\[E\d{3}\]", word_text)
        references = [
            paragraph.text
            for paragraph in word_document.paragraphs
            if paragraph.style.name == "Reference Entry"
        ]
        section = word_document.sections[0]
        word_detail = {
            "path": str(word_path.relative_to(paths.root)),
            "inline_figures": len(word_document.inline_shapes),
            "captions": len(captions),
            "references": len(references),
            "internal_evidence_tokens": len(internal_evidence_tokens),
            "page_width_cm": round(section.page_width.cm, 1),
            "page_height_cm": round(section.page_height.cm, 1),
        }
        word_pass = (
            len(word_document.inline_shapes) == len(figures)
            and len(captions) == len(figures)
            and not internal_evidence_tokens
            and word_detail["page_width_cm"] == 21.0
            and word_detail["page_height_cm"] == 29.7
        )
    checks.append(_check("academic_word_report", word_pass, word_detail))

    network_manifests = list(paths.analyses.glob("network_*_manifest.json"))
    network_errors = []
    for path in network_manifests:
        data = read_json(path)
        if not data.get("empty") and (
            "candidate_pool" not in data
            or "full_candidate_edge_count" not in data
            and "full_within_corpus_edges" not in data
        ):
            network_errors.append(path.name)
    checks.append(
        _check(
            "network_selection_disclosure",
            len(network_manifests) >= 5 and not network_errors,
            {"manifests": len(network_manifests), "incomplete": network_errors},
        )
    )

    export_manifest_path = paths.analyses / "exports" / "export_manifest.json"
    export_detail: Any = "<missing>"
    export_pass = False
    if export_manifest_path.exists():
        export_detail = read_json(export_manifest_path)
        export_pass = (
            export_detail.get("bibliometrix", {}).get("rows") == table_rows.get("works")
            and len(export_detail.get("vosviewer", [])) >= 5
        )
    checks.append(_check("tool_interoperability", export_pass, export_detail))

    scale_path = paths.audit / "scale_plan.json"
    scale_detail = read_json(scale_path) if scale_path.exists() else "<missing>"
    benchmark_path = paths.audit / "scale_benchmark_1m.json"
    benchmark_detail = read_json(benchmark_path) if benchmark_path.exists() else "<missing>"
    checks.append(
        _check(
            "large_scale_policy",
            isinstance(scale_detail, dict)
            and scale_detail.get("policy", {}).get("descriptive_statistics")
            == "all canonical records; never sampled"
            and isinstance(benchmark_detail, dict)
            and benchmark_detail.get("passed") is True
            and benchmark_detail.get("documents", 0) >= 1_000_000,
            {"project_policy": scale_detail, "synthetic_benchmark": benchmark_detail},
        )
    )

    evidence = _load_evidence(paths)
    manuscript_path = paths.report / "manuscript.md"
    manuscript = manuscript_path.read_text(encoding="utf-8")
    generation = read_json(paths.report / "generation_manifest.json")
    validation = validate_manuscript(
        manuscript,
        evidence,
        strict_structure=generation.get("model") != "deterministic-template",
    )
    checks.append(_check("manuscript_evidence_validation", validation["valid"], validation))
    stages = paths.report / "generation_stages"
    stage_files = list(stages.glob("*")) if stages.exists() else []
    specialist_reviews = list(stages.glob("03_review_*_*.md")) if stages.exists() else []
    revised_sections = list(stages.glob("04_revision_*_*.md")) if stages.exists() else []
    quality_path = paths.report / "manuscript_quality.json"
    quality_profile = read_json(quality_path) if quality_path.exists() else {}
    staged_generation_pass = (
        generation.get("pipeline") == "staged-bibliometric-writing-v2"
        and (stages / "01_editorial_plan.md").exists()
        and len(specialist_reviews) >= 3
        and len(revised_sections) >= len(SECTION_CONTRACTS)
        and quality_profile.get("passed") is True
    )
    checks.append(
        _check(
            "staged_journal_ready_generation",
            staged_generation_pass,
            {
                "pipeline": generation.get("pipeline"),
                "stage_artifacts": len(stage_files),
                "specialist_reviews": len(specialist_reviews),
                "revised_sections": len(revised_sections),
                "quality_profile": quality_profile,
            },
        )
    )

    graph = evidence.graph
    used_ids = validation.get("used_evidence_ids", [])
    missing_paths = [
        evidence_id
        for evidence_id in used_ids
        if "raw_snapshot" not in graph
        or evidence_id not in graph
        or not nx.has_path(graph, "raw_snapshot", evidence_id)
    ]
    claim_nodes = [node for node, data in graph.nodes(data=True) if data.get("kind") == "claim"]
    checks.append(
        _check(
            "evidence_graph_traceability",
            not missing_paths and bool(claim_nodes),
            {
                "evidence_items": len(evidence.items),
                "used_evidence": len(used_ids),
                "claim_nodes": len(claim_nodes),
                "missing_raw_to_evidence_paths": missing_paths,
            },
        )
    )

    model_pass = not require_llm_model or generation.get("model") == require_llm_model
    checks.append(
        _check(
            "generation_model",
            model_pass,
            {
                "actual": generation.get("model"),
                "required": require_llm_model,
            },
        )
    )

    result = {
        "project_root": str(paths.root),
        "checked_at": datetime.now(UTC),
        "passed": all(item["passed"] for item in checks),
        "passed_checks": sum(item["passed"] for item in checks),
        "total_checks": len(checks),
        "checks": checks,
    }
    write_json(paths.audit / "acceptance_report.json", result)
    lines = [
        "# CiteWeave 验收报告",
        "",
        f"- 总体结果：{'通过' if result['passed'] else '未通过'}",
        f"- 检查项：{result['passed_checks']}/{result['total_checks']}",
        f"- 项目：`{paths.root}`",
        "",
        "| 检查项 | 结果 |",
        "|---|---|",
        *[f"| {item['name']} | {'通过' if item['passed'] else '失败'} |" for item in checks],
        "",
        "机器可读细节见 `acceptance_report.json`。",
    ]
    (paths.audit / "ACCEPTANCE.md").write_text("\n".join(lines), encoding="utf-8")
    _package_manifest(paths)
    return result
