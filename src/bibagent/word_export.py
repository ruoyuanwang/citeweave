from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from itertools import pairwise
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree

from .figure_catalog import (
    figure_anchor,
    figure_caption,
    figure_reading_expansion,
    figure_reading_guide,
    figure_section,
    order_figures,
)
from .io import read_json
from .models import ProjectPaths

EVIDENCE_RE = re.compile(r"(\[E\d{3}(?:\]\[E\d{3})*\])")
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[E\d{3}(?:\]\[E\d{3})*\])")
REFERENCE_LINE_RE = re.compile(r"^\s*\d+\.\s+(.+?)\s*$")
TRAILING_EVIDENCE_RE = re.compile(r"\s+\[(E\d{3})\]\s*$")
CHINESE_BODY_FONT = "宋体"
CHINESE_HEADING_FONT = "Microsoft YaHei"


@dataclass(frozen=True)
class WordFigure:
    filename: str
    caption: str
    evidence: str
    anchor: str


@dataclass(frozen=True)
class WordReference:
    text: str
    evidence_id: str | None


# Captions are numbered by first occurrence in the paper, not by the HTML
# gallery order.  This makes the DOCX read like a conventional journal article.
DEFAULT_FIGURES = (
    WordFigure(
        "annual_publications.png",
        "年度科学产出趋势",
        "E029",
        "语料库包含269篇唯一文献",
    ),
    WordFigure(
        "document_types.png",
        "文献类型构成",
        "E033",
        "语料库包含269篇唯一文献",
    ),
    WordFigure(
        "top_sources.png",
        "主要来源期刊发文量",
        "E030",
        "在具有来源名称的256篇记录中",
    ),
    WordFigure(
        "bradford_sources.png",
        "Bradford来源分区",
        "E036",
        "在具有来源名称的256篇记录中",
    ),
    WordFigure(
        "top_authors.png",
        "高产作者分布",
        "E031",
        "作者产出方面",
    ),
    WordFigure(
        "top_institutions.png",
        "高产机构分布",
        "E032",
        "机构层面",
    ),
    WordFigure(
        "network_coauthorship.png",
        "作者合作网络",
        "E039",
        "合著候选网络包含122个节点",
    ),
    WordFigure(
        "network_institution_collaboration.png",
        "机构合作网络",
        "E040",
        "合著候选网络包含122个节点",
    ),
    WordFigure(
        "citation_distribution.png",
        "文献被引次数分布",
        "E034",
        "引文分布高度偏斜",
    ),
    WordFigure(
        "top_cited_documents.png",
        "高被引文献",
        "E035",
        "引文分布高度偏斜",
    ),
    WordFigure(
        "keyword_trends.png",
        "关键词年度演化",
        "E037",
        "关键词共现网络包含154个节点",
    ),
    WordFigure(
        "network_keyword_cooccurrence.png",
        "关键词共现网络",
        "E041",
        "关键词共现网络包含154个节点",
    ),
    WordFigure(
        "thematic_map.png",
        "主题图谱",
        "E044",
        "关键词共现网络包含154个节点",
    ),
    WordFigure(
        "three_field_map.png",
        "作者—来源—关键词三字段关系",
        "E038",
        "关键词共现网络包含154个节点",
    ),
    WordFigure(
        "network_cocitation.png",
        "参考文献共被引网络",
        "E042",
        "共被引候选网络包含533个节点",
    ),
    WordFigure(
        "network_bibliographic_coupling.png",
        "文献耦合网络",
        "E043",
        "共被引候选网络包含533个节点",
    ),
)

# A figure may be discussed in more than one results subsection. Insert it only
# in the subsection where that analysis is reported, so cross-references in an
# earlier subsection do not pull the figure out of its conventional position.
FIGURE_RESULT_SECTION_PREFIXES = {
    1: "3.1",
    2: "3.1",
    3: "3.1",
    4: "3.1",
    5: "3.2",
    6: "3.2",
    7: "3.2",
    8: "3.2",
    9: "3.3",
    10: "3.3",
    11: "3.4",
    12: "3.4",
    13: "3.4",
    14: "3.4",
    15: "3.5",
    16: "3.5",
}


def _set_run_font(
    run: Any,
    *,
    latin: str = "Times New Roman",
    east_asia: str = "宋体",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    if bold and east_asia == CHINESE_BODY_FONT:
        # SimSun has no native bold face on many Windows installations.  Use a
        # family with a real bold font file instead of Word's synthetic stroke.
        east_asia = CHINESE_HEADING_FONT
    run.font.name = latin
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:cs"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_style_font(style: Any, *, size: float, east_asia: str, bold: bool = False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    run_properties = style.element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:hint"), "eastAsia")
    language = run_properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run_properties.append(language)
    language.set(qn("w:val"), "en-US")
    language.set(qn("w:eastAsia"), "zh-CN")


def _page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def _keep_with_next(paragraph: Any, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value
    paragraph.paragraph_format.keep_together = True


def _configure_document(document: Document, title: str) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    _set_style_font(normal, size=10.5, east_asia=CHINESE_BODY_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    title_style = document.styles["Title"]
    _set_style_font(
        title_style,
        size=17,
        east_asia=CHINESE_HEADING_FONT,
        bold=True,
    )
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(18)
    title_style.paragraph_format.space_after = Pt(10)
    title_style.paragraph_format.keep_with_next = True

    subtitle_style = document.styles["Subtitle"]
    _set_style_font(subtitle_style, size=10, east_asia="宋体")
    subtitle_style.font.color.rgb = RGBColor(100, 116, 139)
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(18)
    subtitle_style.paragraph_format.keep_with_next = True

    custom_styles = {
        "Abstract Body": (10, "宋体"),
        "Keywords": (10.5, "宋体"),
        "Figure Image": (10, "宋体"),
        "Reference Entry": (9, "宋体"),
    }
    for name, (size, east_asia) in custom_styles.items():
        if name not in document.styles:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = document.styles[name]
        _set_style_font(style, size=size, east_asia=east_asia)
    abstract_style = document.styles["Abstract Body"]
    abstract_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_style.paragraph_format.first_line_indent = Pt(0)
    abstract_style.paragraph_format.space_after = Pt(5)
    abstract_style.paragraph_format.line_spacing = 1.25
    keywords_style = document.styles["Keywords"]
    keywords_style.paragraph_format.first_line_indent = Pt(0)
    keywords_style.paragraph_format.space_before = Pt(4)
    keywords_style.paragraph_format.space_after = Pt(12)
    figure_style = document.styles["Figure Image"]
    figure_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_style.paragraph_format.first_line_indent = Pt(0)
    figure_style.paragraph_format.space_before = Pt(8)
    figure_style.paragraph_format.space_after = Pt(0)
    figure_style.paragraph_format.keep_with_next = True
    figure_style.paragraph_format.keep_together = True
    reference_style = document.styles["Reference Entry"]
    reference_style.paragraph_format.left_indent = Pt(21)
    reference_style.paragraph_format.first_line_indent = Pt(-21)
    reference_style.paragraph_format.space_after = Pt(3)
    reference_style.paragraph_format.line_spacing = 1.05

    heading_tokens = {
        "Heading 1": (14, 14, 8),
        "Heading 2": (12, 12, 6),
        "Heading 3": (11, 8, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = document.styles[name]
        _set_style_font(
            style,
            size=size,
            east_asia=CHINESE_HEADING_FONT,
            bold=True,
        )
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.first_line_indent = Pt(0)

    caption = document.styles["Caption"]
    _set_style_font(caption, size=9, east_asia="宋体")
    caption.font.color.rgb = RGBColor(31, 45, 61)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.keep_together = True

    if "Evidence Note" not in document.styles:
        evidence_note = document.styles.add_style("Evidence Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        evidence_note = document.styles["Evidence Note"]
    _set_style_font(evidence_note, size=8, east_asia="宋体")
    evidence_note.font.color.rgb = RGBColor(98, 112, 132)
    evidence_note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    evidence_note.paragraph_format.space_after = Pt(8)
    evidence_note.paragraph_format.first_line_indent = Pt(0)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("BibAgent 文献计量研究报告")
    _set_run_font(run, east_asia="宋体", size=8.5, color="64748B")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    _page_field(footer)
    for run in footer.runs:
        _set_run_font(run, east_asia="宋体", size=8.5, color="64748B")

    document.core_properties.title = title
    document.core_properties.subject = "Evidence-first bibliometric research report"
    document.core_properties.author = "BibAgent"
    document.core_properties.keywords = "文献计量; 科学知识图谱; Crossref"


def _add_title_block(document: Document, title: str) -> None:
    display_title = re.sub(
        r"^bibliometric主题",
        "Bibliometric 主题",
        title,
        flags=re.IGNORECASE,
    )
    display_title = re.sub("crossref", "Crossref", display_title, flags=re.IGNORECASE)
    display_title = display_title.replace("：基于", "：\n基于", 1)
    title_paragraph = document.add_paragraph(style="Title")
    title_run = title_paragraph.add_run(display_title)
    _set_run_font(
        title_run,
        east_asia=CHINESE_HEADING_FONT,
        size=17,
        bold=True,
    )
    document.add_paragraph(
        "基于可追溯元数据与证据图谱的可审计研究稿",
        style="Subtitle",
    )


def _format_citation(numbers: list[int]) -> str:
    unique = sorted(set(numbers))
    if not unique:
        return ""
    groups: list[str] = []
    start = previous = unique[0]
    for number in unique[1:]:
        if number == previous + 1:
            previous = number
            continue
        groups.append(
            f"{start}–{previous}" if previous - start >= 2 else (
                f"{start},{previous}" if previous != start else str(start)
            )
        )
        start = previous = number
    groups.append(
        f"{start}–{previous}" if previous - start >= 2 else (
            f"{start},{previous}" if previous != start else str(start)
        )
    )
    return f"[{','.join(groups)}]"


def _add_inline_text(
    paragraph: Any,
    text: str,
    *,
    citation_numbers: dict[str, int] | None = None,
    show_citations: bool = True,
    size: float = 10.5,
) -> None:
    citation_numbers = citation_numbers or {}
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            _set_run_font(run, size=size)
        token = match.group()
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=size, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size, italic=True)
        elif show_citations:
            evidence_ids = re.findall(r"E\d{3}", token)
            rendered = _format_citation(
                [
                    citation_numbers[evidence_id]
                    for evidence_id in evidence_ids
                    if evidence_id in citation_numbers
                ]
            )
            if rendered:
                run = paragraph.add_run(rendered)
                _set_run_font(run, size=8.5, color="111827")
                run.font.superscript = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_font(run, size=size)


def _add_body_paragraph(
    document: Document,
    text: str,
    *,
    abstract: bool = False,
    citation_numbers: dict[str, int] | None = None,
) -> Any:
    paragraph = document.add_paragraph(style="Abstract Body" if abstract else "Normal")
    _add_inline_text(
        paragraph,
        text,
        citation_numbers=citation_numbers,
        show_citations=not abstract,
        size=10 if abstract else 10.5,
    )
    return paragraph


def _add_heading(document: Document, text: str, level: int) -> Any:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.first_line_indent = Pt(0)
    return paragraph


def _add_figure(
    document: Document,
    project_paths: ProjectPaths,
    figure: WordFigure,
    number: int,
) -> None:
    path = project_paths.figures / figure.filename
    if not path.exists():
        raise FileNotFoundError(f"Figure is missing: {path}")
    image_paragraph = document.add_paragraph(style="Figure Image")
    _keep_with_next(image_paragraph)
    shape = image_paragraph.add_run().add_picture(str(path), width=Inches(6.1))
    shape._inline.docPr.set("descr", f"图{number} {figure.caption}")

    caption = document.add_paragraph(style="Caption")
    caption_run = caption.add_run(f"图 {number}  {figure.caption}")
    _set_run_font(caption_run, east_asia="宋体", size=9)
    caption.paragraph_format.keep_with_next = True

    note = document.add_paragraph(style="Evidence Note")
    note_run = note.add_run("注：基于规范化元数据与确定性分析结果绘制。")
    _set_run_font(note_run, east_asia="宋体", size=8, color="627084")
    note.paragraph_format.keep_with_next = True

    interpretation = document.add_paragraph(style="Normal")
    lead = interpretation.add_run(f"图{number}直观解读。")
    _set_run_font(
        lead,
        east_asia=CHINESE_HEADING_FONT,
        size=10.5,
        bold=True,
    )
    guide = interpretation.add_run(figure_reading_guide(Path(figure.filename).stem))
    _set_run_font(guide, size=10.5)

    expanded_guide = figure_reading_expansion(Path(figure.filename).stem)
    if expanded_guide:
        expansion = document.add_paragraph(style="Normal")
        _add_inline_text(expansion, expanded_guide, show_citations=False, size=10.5)


def _add_keywords(document: Document, text_value: str) -> None:
    paragraph = document.add_paragraph(style="Keywords")
    label = paragraph.add_run("关键词：")
    _set_run_font(
        label,
        east_asia=CHINESE_HEADING_FONT,
        size=10.5,
        bold=True,
    )
    text = paragraph.add_run(text_value)
    _set_run_font(text, size=10.5)


def _add_reference(document: Document, number: int, text: str) -> None:
    paragraph = document.add_paragraph(style="Reference Entry")
    number_run = paragraph.add_run(f"[{number}] ")
    _set_run_font(number_run, size=9)
    _add_inline_text(paragraph, text, show_citations=False, size=9)


def _prepare_references(
    lines: list[str],
) -> tuple[list[WordReference], dict[str, int]]:
    references: list[WordReference] = []
    in_references = False
    for raw_line in lines:
        line = raw_line.strip()
        if line == "## 参考文献":
            in_references = True
            continue
        if not in_references:
            continue
        match = REFERENCE_LINE_RE.match(line)
        if not match:
            continue
        text = match.group(1)
        evidence_match = TRAILING_EVIDENCE_RE.search(text)
        evidence_id = evidence_match.group(1) if evidence_match else None
        if evidence_match:
            text = text[: evidence_match.start()].rstrip()
        references.append(WordReference(text=text, evidence_id=evidence_id))

    by_evidence = {
        reference.evidence_id: reference
        for reference in references
        if reference.evidence_id
    }
    first_citation_order: list[str] = []
    in_abstract = False
    for raw_line in lines:
        line = raw_line.strip()
        if line == "## 参考文献":
            break
        if line.startswith("## "):
            in_abstract = line[3:].strip() == "摘要"
            continue
        if in_abstract:
            continue
        for evidence_id in re.findall(r"E\d{3}", line):
            if evidence_id in by_evidence and evidence_id not in first_citation_order:
                first_citation_order.append(evidence_id)

    ordered = [by_evidence[evidence_id] for evidence_id in first_citation_order]
    ordered.extend(reference for reference in references if reference.evidence_id is None)
    citation_numbers = {
        reference.evidence_id: number
        for number, reference in enumerate(ordered, start=1)
        if reference.evidence_id
    }
    return ordered, citation_numbers


def normalize_with_microsoft_word(path: Path, *, timeout_seconds: int = 120) -> Path:
    """Round-trip a DOCX through desktop Word for maximum Windows compatibility.

    ``python-docx`` writes valid OOXML, but desktop Word may defer part of its
    layout normalization until the first save.  This optional Windows-only step
    performs that save without changing the publication content.
    """
    resolved = path.resolve()
    if os.name != "nt":
        raise RuntimeError("Microsoft Word normalization is only available on Windows.")
    powershell = shutil.which("powershell.exe") or shutil.which("powershell")
    if not powershell:
        raise RuntimeError("PowerShell is required for Microsoft Word normalization.")

    script = r"""
$ErrorActionPreference = "Stop"
$word = $null
$document = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $document = $word.Documents.Open($env:BIBAGENT_WORD_INPUT, $false, $true)
    $document.SaveAs2($env:BIBAGENT_WORD_OUTPUT, 16)
}
finally {
    if ($document -ne $null) { $document.Close($false) }
    if ($word -ne $null) { $word.Quit() }
    if ($document -ne $null) {
        [System.Runtime.InteropServices.Marshal]::ReleaseComObject($document) | Out-Null
    }
    if ($word -ne $null) {
        [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
    }
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}
"""

    def native_save(source: Path, destination: Path) -> None:
        destination.unlink(missing_ok=True)
        environment = os.environ.copy()
        environment["BIBAGENT_WORD_INPUT"] = str(source)
        environment["BIBAGENT_WORD_OUTPUT"] = str(destination)
        try:
            result = subprocess.run(
                [
                    powershell,
                    "-NoProfile",
                    "-NonInteractive",
                    "-Command",
                    script,
                ],
                check=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                env=environment,
                timeout=timeout_seconds,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
            )
        except subprocess.TimeoutExpired as exc:
            destination.unlink(missing_ok=True)
            raise RuntimeError("Microsoft Word normalization timed out.") from exc
        if result.returncode != 0 or not destination.exists():
            destination.unlink(missing_ok=True)
            detail = (result.stderr or result.stdout).strip()
            raise RuntimeError(f"Microsoft Word normalization failed: {detail}")

    first_save = resolved.with_name(f".{resolved.stem}.word-normalized.docx")
    native_save(resolved, first_save)
    os.replace(first_save, resolved)
    _patch_native_word_styles(resolved)
    second_save = resolved.with_name(f".{resolved.stem}.word-finalized.docx")
    native_save(resolved, second_save)
    os.replace(second_save, resolved)
    return resolved


def _patch_native_word_styles(path: Path) -> None:
    """Set Word theme fonts and remove built-in title decoration.

    Desktop Word rewrites its built-in heading styles to theme references during
    a native save.  Updating the theme keeps those styles Word-native while
    making Chinese font selection deterministic.
    """
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    drawing_namespace = "http://schemas.openxmlformats.org/drawingml/2006/main"

    def word_tag(local_name: str) -> str:
        return f"{{{word_namespace}}}{local_name}"

    def drawing_tag(local_name: str) -> str:
        return f"{{{drawing_namespace}}}{local_name}"

    temporary = path.with_name(f".{path.stem}.style-patched.docx")
    temporary.unlink(missing_ok=True)
    with zipfile.ZipFile(path, "r") as source:
        styles_xml = source.read("word/styles.xml")
        styles_root = etree.fromstring(styles_xml)
        for style in styles_root.findall(word_tag("style")):
            name_element = style.find(word_tag("name"))
            if name_element is None:
                continue
            style_name = str(name_element.get(word_tag("val"), "")).lower()
            if style_name == "title":
                paragraph_properties = style.find(word_tag("pPr"))
                if paragraph_properties is not None:
                    borders = paragraph_properties.find(word_tag("pBdr"))
                    if borders is not None:
                        paragraph_properties.remove(borders)
        patched_styles = etree.tostring(
            styles_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )
        theme_root = etree.fromstring(source.read("word/theme/theme1.xml"))
        font_scheme = theme_root.find(
            f".//{drawing_tag('themeElements')}/{drawing_tag('fontScheme')}"
        )
        if font_scheme is None:
            raise RuntimeError("Word theme does not contain a font scheme.")
        for family_name, east_asia_font in (
            ("majorFont", CHINESE_HEADING_FONT),
            ("minorFont", CHINESE_BODY_FONT),
        ):
            family = font_scheme.find(drawing_tag(family_name))
            if family is None:
                continue
            latin = family.find(drawing_tag("latin"))
            east_asia = family.find(drawing_tag("ea"))
            if latin is not None:
                latin.set("typeface", "Times New Roman")
            if east_asia is not None:
                east_asia.set("typeface", east_asia_font)
            for script_font in family.findall(drawing_tag("font")):
                if script_font.get("script") in {"Hans", "Hant"}:
                    script_font.set("typeface", east_asia_font)
        patched_theme = etree.tostring(
            theme_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )
        with zipfile.ZipFile(temporary, "w") as target:
            for item in source.infolist():
                if item.filename == "word/styles.xml":
                    payload = patched_styles
                elif item.filename == "word/theme/theme1.xml":
                    payload = patched_theme
                else:
                    payload = source.read(item.filename)
                target.writestr(item, payload)
    os.replace(temporary, path)


def export_word_report(
    project_root: Path,
    output_path: Path | None = None,
    *,
    figures: tuple[WordFigure, ...] | None = None,
    native_word: bool = False,
) -> Path:
    """Export an inline-figure academic DOCX from a completed BibAgent project.

    Design basis: ``narrative_proposal``. Named ``academic_article`` overrides:
    A4 paper, 2.54 cm margins, Song/Times typography, black heading hierarchy,
    10.5 pt body at 1.5 lines, and an editorial title block without a cover page.
    """
    paths = ProjectPaths(project_root)
    manuscript_path = paths.report / "manuscript.md"
    if not manuscript_path.exists():
        raise FileNotFoundError(f"Manuscript is missing: {manuscript_path}")
    output = output_path or paths.report / "manuscript.docx"
    output.parent.mkdir(parents=True, exist_ok=True)

    lines = manuscript_path.read_text(encoding="utf-8").splitlines()
    section_corpus: dict[str, list[str]] = {}
    corpus_section = ""
    for raw_line in lines:
        stripped = raw_line.strip()
        if stripped.startswith("### "):
            corpus_section = stripped[4:7]
        elif stripped.startswith("## "):
            corpus_section = ""
        elif corpus_section:
            section_corpus.setdefault(corpus_section, []).append(stripped)
    references, citation_numbers = _prepare_references(lines)
    evidence_map: dict[str, str] = {}
    evidence_path = paths.evidence / "evidence_items.json"
    if evidence_path.exists():
        for item in json.loads(evidence_path.read_text(encoding="utf-8")):
            if item.get("claim_type") != "figure":
                continue
            artifact = Path(str(item.get("artifact_path", ""))).stem
            if artifact:
                evidence_map[artifact] = item["evidence_id"]
    if figures is None:
        manifest_path = paths.figures / "figure_manifest.json"
        if manifest_path.exists():
            manifest = read_json(manifest_path)
            records = manifest.get("figures", []) if isinstance(manifest, dict) else manifest
            figures = tuple(
                WordFigure(
                    f"{item['name']}.png",
                    figure_caption(item["name"]),
                    evidence_map.get(item["name"], ""),
                    figure_anchor(item["name"]),
                )
                for item in order_figures(records)
                if (paths.figures / f"{item['name']}.png").exists()
            )
        else:
            figures = DEFAULT_FIGURES
    else:
        figures = tuple(
            WordFigure(
                figure.filename,
                figure.caption,
                evidence_map.get(Path(figure.filename).stem, figure.evidence),
                figure.anchor,
            )
            for figure in figures
        )
    title = next(
        (line[2:].strip() for line in lines if line.startswith("# ")),
        "文献计量研究报告",
    )
    document = Document()
    _configure_document(document, title)
    _add_title_block(document, title)

    in_abstract = False
    in_references = False
    skip_chart_index = False
    keywords_added = False
    keyword_text = ""
    current_level_two = ""
    inserted_figures: set[int] = set()
    trust_figure_evidence = evidence_path.exists()
    anchored_figures = {
        number
        for number, figure in enumerate(figures or (), start=1)
        if figure.anchor
        and figure.anchor
        in "\n".join(
            section_corpus.get(figure_section(Path(figure.filename).stem), [])
        )
    }

    def insert_section_figures(section_prefix: str) -> None:
        for number, figure in enumerate(figures or (), start=1):
            if number in inserted_figures:
                continue
            if figure_section(Path(figure.filename).stem) == section_prefix:
                _add_figure(document, paths, figure, number)
                inserted_figures.add(number)

    def pending_section_figures() -> list[tuple[int, WordFigure]]:
        section_prefix = current_level_two[:3]
        return [
            (number, figure)
            for number, figure in enumerate(figures or (), start=1)
            if number not in inserted_figures
            and figure_section(Path(figure.filename).stem) == section_prefix
        ]

    def insert_figures_for_text(text: str) -> None:
        """Place figures before the first paragraph that substantively explains them.

        Evidence IDs and catalog anchors take precedence over visible figure
        numbers.  This prevents an early cross-reference—or a stale number in a
        legacy manuscript—from pulling a figure away from its analysis.
        """
        for number, figure in pending_section_figures():
            explicit = f"[[FIGURE:{Path(figure.filename).stem}]]" in text
            anchored = bool(figure.anchor and figure.anchor in text)
            evidenced = bool(figure.evidence and f"[{figure.evidence}]" in text)
            numbered = (
                not trust_figure_evidence
                and re.search(rf"图\s*{number}(?!\d)", text) is not None
            )
            matched = (
                explicit
                or anchored
                or (
                    number not in anchored_figures
                    and (evidenced or numbered)
                )
            )
            if matched:
                _add_figure(document, paths, figure, number)
                inserted_figures.add(number)

    def split_at_figure_anchors(text: str) -> list[str]:
        """Split a mixed analysis paragraph where a second figure topic begins."""
        split_positions = {
            text.index(figure.anchor)
            for _, figure in pending_section_figures()
            if figure.anchor and figure.anchor in text and text.index(figure.anchor) > 0
        }
        if not split_positions:
            return [text]
        boundaries = [0, *sorted(split_positions), len(text)]
        return [
            text[start:end].strip()
            for start, end in pairwise(boundaries)
            if text[start:end].strip()
        ]

    for raw_line in lines:
        line = raw_line.strip()
        if not line or line.startswith("# "):
            continue
        if line == "## 图表解读清单":
            skip_chart_index = True
            continue
        if line == "## 参考文献":
            skip_chart_index = False
            in_references = True
            _add_heading(document, "参考文献", 1)
            for number, reference in enumerate(references, start=1):
                _add_reference(document, number, reference.text)
            continue
        if skip_chart_index or line.startswith("|"):
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if current_level_two.startswith("3."):
                insert_section_figures(current_level_two[:3])
            if in_abstract and not keywords_added:
                _add_keywords(document, keyword_text or "文献计量；科学知识图谱")
                keywords_added = True
            in_abstract = heading == "摘要"
            in_references = False
            paragraph = _add_heading(document, heading, 1)
            if in_abstract:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("#### "):
            _add_heading(document, line[5:].strip(), 3)
            continue
        if line.startswith("### "):
            if current_level_two.startswith("3."):
                insert_section_figures(current_level_two[:3])
            current_level_two = line[4:].strip()
            _add_heading(document, current_level_two, 2)
            continue
        if line.startswith("**关键词**"):
            keyword_text = line.split("：", 1)[-1].strip("* ")
            if not keywords_added:
                _add_keywords(document, keyword_text)
                keywords_added = True
            continue
        if in_references:
            continue

        for segment in split_at_figure_anchors(line):
            insert_figures_for_text(segment)
            clean_line = re.sub(
                r"\[\[FIGURE:[A-Za-z0-9_-]+\]\]",
                "",
                segment,
            ).strip()
            if not clean_line:
                continue
            _add_body_paragraph(
                document,
                clean_line,
                abstract=in_abstract,
                citation_numbers=citation_numbers,
            )

    if current_level_two.startswith("3."):
        insert_section_figures(current_level_two[:3])
    if len(inserted_figures) != len(figures):
        missing = sorted(set(range(1, len(figures) + 1)) - inserted_figures)
        raise ValueError(
            "Every figure must map to a results subsection; "
            f"missing figure numbers: {missing}"
        )
    document.save(output)
    if native_word:
        normalize_with_microsoft_word(output)
    return output.resolve()
